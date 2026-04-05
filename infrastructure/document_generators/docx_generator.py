import logging
import uuid
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
""" from docx.oxml.ns import qn
from docx.oxml import OxmlElement """

from domain.exceptions import DocumentGenerationError
from domain.summary import SummaryResult

logger = logging.getLogger(__name__)


def generate_docx(result: SummaryResult, output_dir: Path) -> Path:
    """
    Generate a styled DOCX from a SummaryResult.
    Returns the path to the generated file.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    file_path = output_dir / f"summary_{uuid.uuid4().hex[:8]}.docx"

    try:
        _build_docx(result, file_path)
    except Exception as exc:
        logger.exception("DOCX generation failed")
        raise DocumentGenerationError(f"Failed to generate DOCX: {exc}") from exc

    logger.info("DOCX generated at %s", file_path)
    return file_path


def _build_docx(result: SummaryResult, file_path: Path) -> None:
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading("Document Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_color(title.runs[0], RGBColor(0x1A, 0x1A, 0x2E))

    # ── Metadata ───────────────────────────────────────────────────────────
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = meta_para.add_run(
        f"Original: {result.word_count_original} words  |  "
        f"Summary: {result.word_count_summary} words  |  "
        f"Reduction: {_reduction_pct(result)}%"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ── Summary Section ────────────────────────────────────────────────────
    summary_heading = doc.add_heading("Summary", level=2)
    _set_run_color(summary_heading.runs[0], RGBColor(0x16, 0x21, 0x3E))

    body = doc.add_paragraph(result.summary)
    body.style.font.size = Pt(11)

    # ── Key Points Section ─────────────────────────────────────────────────
    if result.key_points:
        doc.add_paragraph()
        kp_heading = doc.add_heading("Key Points", level=2)
        _set_run_color(kp_heading.runs[0], RGBColor(0x16, 0x21, 0x3E))

        for point in result.key_points:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(point)
            run.font.size = Pt(11)

    doc.save(str(file_path))


def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def _reduction_pct(result: SummaryResult) -> int:
    if result.word_count_original == 0:
        return 0
    return round(
        (1 - result.word_count_summary / result.word_count_original) * 100
    )